"""File processing utilities for TaskPilot AI."""
from __future__ import annotations

import base64
import io
import zipfile
from pathlib import Path
from typing import Any

import logging

logger = logging.getLogger(__name__)


class FileProcessor:
    """Process various file types for task execution."""
    
    @staticmethod
    async def process_attachment(mime_type: str, data: str) -> dict[str, Any]:
        """Process an attachment and extract meaningful content."""
        try:
            # Decode base64 data
            if ',' in data:
                data = data.split(',', 1)[1]
            
            file_bytes = base64.b64decode(data)
            
            # Route to appropriate processor based on mime type
            if 'image/' in mime_type:
                return await FileProcessor._process_image(file_bytes, mime_type)
            elif mime_type == 'application/pdf':
                return await FileProcessor._process_pdf(file_bytes)
            elif mime_type in ['text/plain', 'text/csv', 'text/markdown', 'text/html']:
                return await FileProcessor._process_text(file_bytes, mime_type)
            elif mime_type in ['application/json', 'text/json']:
                return await FileProcessor._process_json(file_bytes)
            elif mime_type in ['application/vnd.openxmlformats-officedocument.wordprocessingml.document', 
                              'application/msword']:
                return await FileProcessor._process_docx(file_bytes)
            elif mime_type in ['application/vnd.openxmlformats-officedocument.spreadsheetml.sheet',
                              'application/vnd.ms-excel']:
                return await FileProcessor._process_excel(file_bytes)
            elif mime_type == 'application/zip':
                return await FileProcessor._process_zip(file_bytes)
            else:
                # Try to decode as text as fallback
                try:
                    text_content = file_bytes.decode('utf-8', errors='ignore')[:1000]
                    if text_content.strip():
                        return await FileProcessor._process_text(file_bytes, 'text/plain')
                except Exception:
                    pass
                
                return {
                    "type": "unknown",
                    "mime_type": mime_type,
                    "content": f"Received file of type {mime_type} ({len(file_bytes)} bytes)",
                    "metadata": {"size": len(file_bytes)}
                }
        except Exception as e:
            logger.error(f"Error processing attachment: {e}")
            return {
                "type": "error",
                "content": f"Could not process file: {type(e).__name__}",
                "metadata": {"error": str(e)}
            }
    
    @staticmethod
    async def _process_image(file_bytes: bytes, mime_type: str) -> dict[str, Any]:
        """Process image file - return metadata for Gemini Vision processing."""
        return {
            "type": "image",
            "mime_type": mime_type,
            "content": "Image uploaded - will be analyzed using vision capabilities",
            "base64_data": base64.b64encode(file_bytes).decode('utf-8'),
            "metadata": {"size": len(file_bytes), "format": mime_type.split('/')[-1]}
        }
    
    @staticmethod
    async def _process_pdf(file_bytes: bytes) -> dict[str, Any]:
        """Extract text from PDF file."""
        try:
            import PyPDF2
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
            
            text_parts = []
            for i, page in enumerate(pdf_reader.pages[:50]):  # Limit to first 50 pages
                text = page.extract_text()
                if text.strip():
                    text_parts.append(f"[Page {i+1}]\n{text.strip()}")
            
            full_text = "\n\n".join(text_parts)
            return {
                "type": "pdf",
                "content": full_text[:10000],  # Limit to prevent token overflow
                "metadata": {
                    "pages": len(pdf_reader.pages),
                    "extracted_length": len(full_text)
                }
            }
        except ImportError:
            return {
                "type": "pdf",
                "content": "PDF file received (PyPDF2 not installed for text extraction)",
                "metadata": {"size": len(file_bytes)}
            }
        except Exception as e:
            logger.error(f"PDF processing error: {e}")
            return {
                "type": "pdf",
                "content": f"PDF file received but could not extract text: {type(e).__name__}",
                "metadata": {"size": len(file_bytes), "error": str(e)}
            }
    
    @staticmethod
    async def _process_text(file_bytes: bytes, mime_type: str) -> dict[str, Any]:
        """Process text-based files."""
        try:
            text = file_bytes.decode('utf-8', errors='ignore')
            
            if mime_type == 'text/csv':
                # Try to parse CSV for structured data
                try:
                    import csv
                    reader = csv.DictReader(io.StringIO(text))
                    rows = list(reader)[:100]  # Limit rows
                    
                    return {
                        "type": "csv",
                        "content": text[:5000],
                        "structured_data": rows,
                        "metadata": {
                            "rows": len(rows),
                            "columns": list(rows[0].keys()) if rows else []
                        }
                    }
                except Exception:
                    pass  # Fall back to plain text
            
            return {
                "type": "text",
                "mime_type": mime_type,
                "content": text[:10000],  # Limit size
                "metadata": {"length": len(text)}
            }
        except Exception as e:
            logger.error(f"Text processing error: {e}")
            return {
                "type": "text",
                "content": f"Text file received but could not decode: {type(e).__name__}",
                "metadata": {"size": len(file_bytes)}
            }
    
    @staticmethod
    async def _process_docx(file_bytes: bytes) -> dict[str, Any]:
        """Extract text from DOCX file."""
        try:
            from docx import Document
            doc = Document(io.BytesIO(file_bytes))
            
            paragraphs = []
            for para in doc.paragraphs[:500]:  # Limit paragraphs
                if para.text.strip():
                    paragraphs.append(para.text.strip())
            
            full_text = "\n\n".join(paragraphs)
            return {
                "type": "docx",
                "content": full_text[:10000],
                "metadata": {
                    "paragraphs": len(doc.paragraphs),
                    "extracted_length": len(full_text)
                }
            }
        except ImportError:
            return {
                "type": "docx",
                "content": "Word document received (python-docx not installed for text extraction)",
                "metadata": {"size": len(file_bytes)}
            }
        except Exception as e:
            logger.error(f"DOCX processing error: {e}")
            return {
                "type": "docx",
                "content": f"Word document received but could not extract text: {type(e).__name__}",
                "metadata": {"size": len(file_bytes)}
            }
    
    @staticmethod
    async def _process_zip(file_bytes: bytes) -> dict[str, Any]:
        """Unzip and process contents of ZIP file."""
        try:
            with zipfile.ZipFile(io.BytesIO(file_bytes)) as zf:
                file_list = zf.namelist()[:50]  # Limit number of files
                
                extracted_contents = []
                for filename in file_list:
                    try:
                        # Skip directories and hidden files
                        if filename.endswith('/') or filename.startswith('.'):
                            continue
                        
                        # Limit file size
                        file_info = zf.getinfo(filename)
                        if file_info.file_size > 5_000_000:  # 5MB limit per file
                            extracted_contents.append({
                                "filename": filename,
                                "content": f"File too large ({file_info.file_size} bytes)",
                                "skipped": True
                            })
                            continue
                        
                        content_bytes = zf.read(filename)
                        
                        # Try to decode as text
                        try:
                            content = content_bytes.decode('utf-8', errors='ignore')[:5000]
                            extracted_contents.append({
                                "filename": filename,
                                "content": content,
                                "type": "text",
                                "size": len(content_bytes)
                            })
                        except Exception:
                            extracted_contents.append({
                                "filename": filename,
                                "content": f"Binary file ({len(content_bytes)} bytes)",
                                "type": "binary",
                                "size": len(content_bytes)
                            })
                    except Exception as e:
                        logger.warning(f"Could not extract {filename}: {e}")
                        extracted_contents.append({
                            "filename": filename,
                            "content": f"Error extracting: {type(e).__name__}",
                            "error": True
                        })
                
                return {
                    "type": "zip",
                    "content": f"ZIP archive with {len(file_list)} files",
                    "files": extracted_contents,
                    "metadata": {
                        "total_files": len(zf.namelist()),
                        "processed_files": len(extracted_contents)
                    }
                }
        except Exception as e:
            logger.error(f"ZIP processing error: {e}")
            return {
                "type": "zip",
                "content": f"ZIP file received but could not extract: {type(e).__name__}",
                "metadata": {"size": len(file_bytes), "error": str(e)}
            }
    
    @staticmethod
    async def _process_json(file_bytes: bytes) -> dict[str, Any]:
        """Process JSON file and extract structured data."""
        try:
            import json
            text = file_bytes.decode('utf-8', errors='ignore')
            data = json.loads(text)
            
            # Create a readable summary
            if isinstance(data, dict):
                keys = list(data.keys())[:20]
                summary = f"JSON object with {len(data)} keys: {', '.join(keys)}"
            elif isinstance(data, list):
                summary = f"JSON array with {len(data)} items"
            else:
                summary = f"JSON data: {type(data).__name__}"
            
            # Convert to readable text
            json_text = json.dumps(data, indent=2)[:5000]
            
            return {
                "type": "json",
                "content": f"{summary}\n\n{json_text}",
                "structured_data": data if len(str(data)) < 50000 else None,
                "metadata": {
                    "data_type": type(data).__name__,
                    "size": len(text)
                }
            }
        except Exception as e:
            logger.error(f"JSON processing error: {e}")
            return {
                "type": "json",
                "content": f"JSON file received but could not parse: {type(e).__name__}",
                "metadata": {"size": len(file_bytes)}
            }
    
    @staticmethod
    async def _process_excel(file_bytes: bytes) -> dict[str, Any]:
        """Process Excel file and extract data."""
        try:
            import openpyxl
            workbook = openpyxl.load_workbook(io.BytesIO(file_bytes), read_only=True, data_only=True)
            
            sheets_data = []
            for sheet_name in workbook.sheetnames[:10]:  # Limit to 10 sheets
                sheet = workbook[sheet_name]
                
                # Extract data (limit rows)
                data = []
                for i, row in enumerate(sheet.iter_rows(values_only=True)):
                    if i >= 100:  # Limit to 100 rows per sheet
                        break
                    data.append([str(cell) if cell is not None else "" for cell in row])
                
                if data:
                    sheets_data.append({
                        "sheet": sheet_name,
                        "rows": len(data),
                        "columns": len(data[0]) if data else 0,
                        "data": data[:20]  # First 20 rows only in response
                    })
            
            # Create text summary
            summary_parts = []
            for sheet_info in sheets_data:
                summary_parts.append(f"Sheet '{sheet_info['sheet']}': {sheet_info['rows']} rows × {sheet_info['columns']} columns")
                
                # Include first few rows as preview
                if sheet_info['data']:
                    for i, row in enumerate(sheet_info['data'][:5]):
                        row_preview = " | ".join(row[:10])  # First 10 columns
                        summary_parts.append(f"Row {i+1}: {row_preview}")
            
            content = "\n".join(summary_parts)
            
            return {
                "type": "excel",
                "content": content[:8000],
                "structured_data": sheets_data,
                "metadata": {
                    "sheets": len(sheets_data),
                    "total_rows": sum(s['rows'] for s in sheets_data)
                }
            }
        except ImportError:
            return {
                "type": "excel",
                "content": "Excel file received (openpyxl not installed for data extraction)",
                "metadata": {"size": len(file_bytes)}
            }
        except Exception as e:
            logger.error(f"Excel processing error: {e}")
            return {
                "type": "excel",
                "content": f"Excel file received but could not extract data: {type(e).__name__}",
                "metadata": {"size": len(file_bytes), "error": str(e)}
            }

